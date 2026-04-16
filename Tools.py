from typing import Dict, Any, List, Optional
from pathlib import Path
import json
from docx import Document

# 工具箱，用来注册，提取工具
class ToolExecutor:
    """
    一个工具执行器，负责管理和执行工具。
    """
    def __init__(self):
        self.tools: Dict[str, Dict[str, Any]] = {}

    def registerTool(self, name: str, description: str, func: callable):
        """
        向工具箱中注册一个新工具。
        """
        if name in self.tools:
            print(f"警告:工具 '{name}' 已存在，将被覆盖。")
        self.tools[name] = {"description": description, "func": func}
        print(f"工具 '{name}' 已注册。")

    def getTool(self, name: str) -> callable:
        """
        根据名称获取一个工具的执行函数。
        """
        return self.tools.get(name, {}).get("func")

    def getAvailableTools(self) -> str:
        """
        获取所有可用工具的格式化描述字符串。
        """
        return "\n".join([
            f"- {name}: {info['description']}"
            for name, info in self.tools.items()
        ])

# 读取verilog代码工具
# 输入文件路径，开始读取的行数，结束读取的行数，返回一个字典
def rd_verilog(file_path: str, start: Optional[int] = None, end: Optional[int] = None) -> Dict[str, Any]:
    """
    读取 Verilog/SystemVerilog 源码文件。

    参数:
        file_path: Verilog 文件路径
        start: 起始行号（从 1 开始，可选）
        end: 结束行号（从 1 开始，可选，包含该行）

    返回:
        dict:
        {
            "filename": 文件名,
            "filepath": 绝对路径,
            "source_code": 源码字符串,
            "line_count": 文件总行数,
            "read_ok": 是否读取成功,
            "error": 错误信息
        }
    """
    result = {
        "filename": "",
        "filepath": "",
        "source_code": "",
        "line_count": 0,
        "read_ok": False,
        "error": ""
    }

    try:
        # 基本类型检查
        if not isinstance(file_path, str) or not file_path.strip():
            raise ValueError("file_path 必须是非空字符串。")

        if start is not None and (not isinstance(start, int) or start <= 0):
            raise ValueError("start 必须是正整数或 None。")

        if end is not None and (not isinstance(end, int) or end <= 0):
            raise ValueError("end 必须是正整数或 None。")

        if start is not None and end is not None and start > end:
            raise ValueError("start 不能大于 end。")

        path = Path(file_path).expanduser().resolve()

        result["filename"] = path.name
        result["filepath"] = str(path)

        # 文件存在性检查
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        if not path.is_file():
            raise FileNotFoundError(f"路径不是文件: {path}")

        # 可按需限制后缀
        if path.suffix.lower() not in {".v", ".sv"}:
            raise ValueError(f"不支持的文件类型: {path.suffix}，仅支持 .v 和 .sv 文件。")

        # 尝试多种编码读取
        encodings = ["utf-8", "gbk", "latin-1"]
        file_text = None
        last_error = None

        for enc in encodings:
            try:
                file_text = path.read_text(encoding=enc)
                break
            except UnicodeDecodeError as e:
                last_error = e

        if file_text is None:
            raise UnicodeDecodeError(
                getattr(last_error, "encoding", "unknown"),
                getattr(last_error, "object", b""),
                getattr(last_error, "start", 0),
                getattr(last_error, "end", 1),
                "无法使用 utf-8 / gbk / latin-1 解码文件。"
            )

        lines = file_text.splitlines()
        total_lines = len(lines)
        result["line_count"] = total_lines

        # 若未指定范围，则返回全部源码
        if start is None and end is None:
            selected_code = file_text
        else:
            # 只给了一个参数时的默认行为
            actual_start = start if start is not None else 1
            actual_end = end if end is not None else total_lines

            if actual_start > total_lines:
                raise ValueError(
                    f"start={actual_start} 超出文件总行数 {total_lines}。"
                )

            if actual_end > total_lines:
                actual_end = total_lines

            # Python 切片左闭右开，因此 end 不需要再减 1
            selected_lines = lines[actual_start - 1: actual_end]
            selected_code = "\n".join(selected_lines)

        result["source_code"] = selected_code
        result["read_ok"] = True
        result["error"] = ""

    except Exception as e:
        result["read_ok"] = False
        result["error"] = str(e)

    return result

# 写json工具
class JsonSummaryStore:
    """
    管理 Verilog 摘要 JSON 的工具类。
    支持:
    - add: 新增一个文件条目
    - update: 更新指定文件中的某一个块
    - list: 列出当前已处理文件
    - query: 查询指定文件条目
    """

    VALID_BLOCKS = {
        "code_function_descriptions",
        "code_indices",
        "reflection_feedback",
    }

    MAX_FILE_COUNT = 25

    def __init__(self, json_db_path: str):
        self.db_path = Path(json_db_path).expanduser().resolve()

    # ========== 基础读写 ==========
    def _init_db(self) -> Dict[str, Any]:
        return {
            "files": []
        }

    def _load(self) -> Dict[str, Any]:
        if not self.db_path.exists():
            return self._init_db()

        try:
            text = self.db_path.read_text(encoding="utf-8")
            if not text:
                data = self._init_db()
                self._save(data)
                return data

            data = json.loads(text)

            if not isinstance(data, dict):
                raise ValueError("JSON 根节点必须是字典。")
            if "files" not in data or not isinstance(data["files"], list):
                raise ValueError("JSON 中必须包含 'files' 列表。")

            return data
        except Exception as e:
            raise ValueError(f"读取 JSON 数据库失败: {e}")

    def _save(self, data: Dict[str, Any]) -> None:
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        self.db_path.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

    # ========== 辅助函数 ==========
    def _normalize_filename(self, filename: str) -> str:
        if not isinstance(filename, str) or not filename.strip():
            raise ValueError("filename 必须是非空字符串。")
        return filename.strip()

    def _find_entry_by_filename(self, data: Dict[str, Any], filename: str) -> Optional[Dict[str, Any]]:
        filename = self._normalize_filename(filename)
        for entry in data["files"]:
            if entry.get("filename") == filename:
                return entry
        return None

    def _generate_next_file_id(self, data: Dict[str, Any]) -> str:
        """
        为新文件生成 file_id: F00 ~ F24
        采用“找第一个未占用编号”的方式。
        """
        used_ids = set()
        for entry in data["files"]:
            fid = entry.get("file_id")
            if isinstance(fid, str):
                used_ids.add(fid)

        for i in range(self.MAX_FILE_COUNT):
            candidate = f"F{i:02d}"
            if candidate not in used_ids:
                return candidate

        raise ValueError(f"文件数量已达到上限 {self.MAX_FILE_COUNT}，无法再分配 file_id。")

    def _validate_file_entry(self, file_entry: Dict[str, Any]) -> Dict[str, Any]:
        """
        校验并规范化新增文件条目。
        add 时 file_id 由系统自动补充。
        """
        if not isinstance(file_entry, dict):
            raise ValueError("file_entry 必须是字典。")

        filename = file_entry.get("filename", "")
        filepath = file_entry.get("filepath", "")

        if not isinstance(filename, str) or not filename.strip():
            raise ValueError("file_entry['filename'] 必须是非空字符串。")
        if not filename.endswith(".v") and not filename.endswith(".sv"):
            raise ValueError("filename 必须是 .v 或 .sv 文件。")

        if not isinstance(filepath, str) or not filepath.strip():
            raise ValueError("file_entry['filepath'] 必须是非空字符串。")

        code_function_descriptions = file_entry.get("code_function_descriptions", [])
        code_indices = file_entry.get("code_indices", [])
        reflection_feedback = file_entry.get("reflection_feedback", [])

        if not isinstance(code_function_descriptions, list):
            raise ValueError("code_function_descriptions 必须是列表。")
        if not isinstance(code_indices, list):
            raise ValueError("code_indices 必须是列表。")
        if not isinstance(reflection_feedback, list):
            raise ValueError("reflection_feedback 必须是列表。")

        normalized = {
            "filename": filename.strip(),
            "filepath": str(Path(filepath).expanduser().resolve()),
            "code_function_descriptions": code_function_descriptions,
            "code_indices": code_indices,
            "reflection_feedback": reflection_feedback,
        }
        return normalized

    # ========== 业务功能 ==========
    def add(self, file_entry: Dict[str, Any]) -> Dict[str, Any]:
        """
        在 JSON 末尾追加一个新的文件条目。
        若 filename 已存在，则报错，不覆盖。
        """
        data = self._load()
        normalized = self._validate_file_entry(file_entry)

        existing = self._find_entry_by_filename(data, normalized["filename"])
        if existing is not None:
            raise ValueError(f"文件 '{normalized['filename']}' 已存在，不能重复 add。")

        file_id = self._generate_next_file_id(data)
        normalized["file_id"] = file_id

        data["files"].append(normalized)
        self._save(data)

        return {
            "success": True,
            "message": f"已新增文件 {normalized['filename']}，分配标识符 {file_id}。",
            "data": normalized
        }

    def update(self, file_name: str, block_name: str, new_content: List[Any]) -> Dict[str, Any]:
        """
        更新指定文件的某一个块。
        block_name 只能是:
        - code_function_descriptions
        - code_indices
        - reflection_feedback

        每次只允许修改一个块。
        """
        if block_name not in self.VALID_BLOCKS:
            raise ValueError(
                f"block_name 非法，必须是 {sorted(self.VALID_BLOCKS)} 之一。"
            )

        if not isinstance(new_content, list):
            raise ValueError("new_content 必须是列表。")

        data = self._load()
        entry = self._find_entry_by_filename(data, file_name)
        if entry is None:
            raise KeyError(f"未找到文件: {file_name}")

        entry[block_name] = new_content
        self._save(data)

        return {
            "success": True,
            "message": f"已更新文件 {file_name} 的 {block_name}。",
            "data": {
                "file_id": entry.get("file_id", ""),
                "filename": entry.get("filename", ""),
                "updated_block": block_name,
                "new_content": entry[block_name]
            }
        }

    def list_files(self) -> Dict[str, Any]:
        """
        列举当前 JSON 中已经处理过的文件。
        默认按 add 顺序返回。
        """
        data = self._load()
        files = []

        for entry in data["files"]:
            files.append({
                "file_id": entry.get("file_id", ""),
                "filename": entry.get("filename", ""),
                "filepath": entry.get("filepath", "")
            })

        return {
            "success": True,
            "message": f"当前已处理 {len(files)} 个文件。",
            "data": files
        }

    def query(self, file_name: str) -> Dict[str, Any]:
        """
        查询指定文件的完整 JSON 片段。
        """
        data = self._load()
        entry = self._find_entry_by_filename(data, file_name)
        if entry is None:
            raise KeyError(f"未找到文件: {file_name}")

        return {
            "success": True,
            "message": f"已查询到文件 {file_name} 的内容。",
            "data": entry
        }


def wr_json(operation: str, json_db_path: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    """
    统一入口函数。

    参数:
        operation:
            - "add"
            - "update"
            - "updata"   # 兼容你的拼写
            - "list"
            - "query"

        json_db_path:
            摘要 JSON 数据库路径

        payload:
            不同操作对应不同字段

    返回:
        dict:
        {
            "success": bool,
            "message": str,
            "data": ...
        }
    """
    try:
        store = JsonSummaryStore(json_db_path)
        op = operation.strip().lower()

        if op == "add":
            file_entry = payload.get("file_entry")
            if file_entry is None:
                raise ValueError("add 操作需要 payload['file_entry']")
            return store.add(file_entry)

        elif op in {"update", "updata"}:
            file_name = payload.get("file_name")
            block_name = payload.get("block_name")
            new_content = payload.get("new_content")

            if file_name is None:
                raise ValueError("update 操作需要 payload['file_name']")
            if block_name is None:
                raise ValueError("update 操作需要 payload['block_name']")
            if new_content is None:
                raise ValueError("update 操作需要 payload['new_content']")

            return store.update(file_name, block_name, new_content)

        elif op == "list":
            return store.list_files()

        elif op == "query":
            file_name = payload.get("file_name")
            if file_name is None:
                raise ValueError("query 操作需要 payload['file_name']")
            return store.query(file_name)

        else:
            raise ValueError("operation 仅支持 add / update / updata / list / query")

    except Exception as e:
        return {
            "success": False,
            "message": f"{type(e).__name__}: {e}",
            "data": None
        }


# =========================
# 写 DOC JSON 工具
# =========================
class DocJsonStore:
    """
    管理 DOC 文档 JSON 映射的工具类。
    目标：让 LLM 以“小节”为粒度增改删查，而不是每次全文重写。

    根结构:
    {
      "doc_name": "firmware.docx",
      "doc_path": "/abs/path/firmware.docx",
      "doc_title": "固件说明文档",
      "sections": [...]
    }

    每个小节结构:
    {
      "section_id": "S001",
      "section_number": "1.1",
      "section_title": "时基初始化",
      "heading_level": 2,
      "content": [
        {"block_id": "B001", "block_type": "paragraph", "text": "..."},
        {"block_id": "B002", "block_type": "code", "text": "...", "language": "c"}
      ],
      "source_refs": [
        {"file_name": "clock.v", "line_start": 1, "line_end": 25}
      ],
      "tags": ["初始化", "时钟"]
    }
    """

    VALID_SECTION_BLOCKS = {
        "section_number",
        "section_title",
        "heading_level",
        "content",
        "source_refs",
        "tags",
    }

    def __init__(self, doc_json_path: str):
        self.db_path = Path(doc_json_path).expanduser().resolve()

    # ========== 基础读写 ==========
    def _init_db(self) -> Dict[str, Any]:
        return {
            "doc_name": "",
            "doc_path": "",
            "doc_title": "",
            "sections": []
        }

    def _load(self) -> Dict[str, Any]:
        if not self.db_path.exists():
            return self._init_db()

        try:
            text = self.db_path.read_text(encoding="utf-8").strip()
            if not text:
                data = self._init_db()
                self._save(data)
                return data

            data = json.loads(text)
            if not isinstance(data, dict):
                raise ValueError("DOC JSON 根节点必须是字典。")
            if "sections" not in data or not isinstance(data["sections"], list):
                raise ValueError("DOC JSON 中必须包含 'sections' 列表。")

            data.setdefault("doc_name", "")
            data.setdefault("doc_path", "")
            data.setdefault("doc_title", "")
            return data
        except Exception as e:
            raise ValueError(f"读取 DOC JSON 数据库失败: {e}")

    def _save(self, data: Dict[str, Any]) -> None:
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        self.db_path.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

    # ========== 辅助函数 ==========
    def _normalize_non_empty_str(self, value: Any, field_name: str) -> str:
        if not isinstance(value, str) or not value.strip():
            raise ValueError(f"{field_name} 必须是非空字符串。")
        return value.strip()

    def _normalize_optional_str(self, value: Any, field_name: str) -> str:
        if value is None:
            return ""
        if not isinstance(value, str):
            raise ValueError(f"{field_name} 必须是字符串或 None。")
        return value.strip()

    def _generate_next_section_id(self, data: Dict[str, Any]) -> str:
        used_ids = set()
        for entry in data["sections"]:
            sid = entry.get("section_id")
            if isinstance(sid, str):
                used_ids.add(sid)

        i = 1
        while True:
            candidate = f"S{i:03d}"
            if candidate not in used_ids:
                return candidate
            i += 1

    def _generate_next_block_id(self, content: List[Dict[str, Any]]) -> str:
        used_ids = set()
        for block in content:
            bid = block.get("block_id")
            if isinstance(bid, str):
                used_ids.add(bid)

        i = 1
        while True:
            candidate = f"B{i:03d}"
            if candidate not in used_ids:
                return candidate
            i += 1

    def _normalize_content_block(self, block: Dict[str, Any], auto_block_id: Optional[str] = None) -> Dict[str, Any]:
        if not isinstance(block, dict):
            raise ValueError("content 中的每个 block 都必须是字典。")

        block_type = self._normalize_non_empty_str(block.get("block_type"), "block_type")
        if block_type not in {"paragraph", "code", "bullet", "table_note"}:
            raise ValueError("block_type 仅支持 paragraph / code / bullet / table_note。")

        text = block.get("text")
        if not isinstance(text, str):
            raise ValueError("content block 的 text 必须是字符串。")

        block_id = block.get("block_id")
        if block_id is None:
            block_id = auto_block_id or ""
        if block_id and not isinstance(block_id, str):
            raise ValueError("block_id 必须是字符串。")

        normalized = {
            "block_id": block_id,
            "block_type": block_type,
            "text": text
        }

        if "language" in block:
            if not isinstance(block["language"], str):
                raise ValueError("language 必须是字符串。")
            normalized["language"] = block["language"]

        return normalized

    def _normalize_source_ref(self, ref: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(ref, dict):
            raise ValueError("source_refs 中的每个元素都必须是字典。")

        file_name = self._normalize_non_empty_str(ref.get("file_name"), "source_refs.file_name")

        line_start = ref.get("line_start")
        line_end = ref.get("line_end")

        if line_start is not None and (not isinstance(line_start, int) or line_start <= 0):
            raise ValueError("source_refs.line_start 必须是正整数或 None。")
        if line_end is not None and (not isinstance(line_end, int) or line_end <= 0):
            raise ValueError("source_refs.line_end 必须是正整数或 None。")
        if line_start is not None and line_end is not None and line_start > line_end:
            raise ValueError("source_refs.line_start 不能大于 source_refs.line_end。")

        normalized = {"file_name": file_name}
        if "file_path" in ref and ref["file_path"] is not None:
            if not isinstance(ref["file_path"], str) or not ref["file_path"].strip():
                raise ValueError("source_refs.file_path 必须是非空字符串。")
            normalized["file_path"] = str(Path(ref["file_path"]).expanduser().resolve())
        if line_start is not None:
            normalized["line_start"] = line_start
        if line_end is not None:
            normalized["line_end"] = line_end
        return normalized

    def _normalize_section_entry(self, section_entry: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(section_entry, dict):
            raise ValueError("section_entry 必须是字典。")

        section_number = self._normalize_optional_str(section_entry.get("section_number"), "section_number")
        section_title = self._normalize_non_empty_str(section_entry.get("section_title"), "section_title")

        heading_level = section_entry.get("heading_level", 2)
        if not isinstance(heading_level, int) or heading_level <= 0:
            raise ValueError("heading_level 必须是正整数。")

        content = section_entry.get("content", [])
        if not isinstance(content, list):
            raise ValueError("content 必须是列表。")

        normalized_content = []
        for block in content:
            normalized_block = self._normalize_content_block(
                block,
                auto_block_id=self._generate_next_block_id(normalized_content)
            )
            if not normalized_block["block_id"]:
                normalized_block["block_id"] = self._generate_next_block_id(normalized_content)
            normalized_content.append(normalized_block)

        source_refs = section_entry.get("source_refs", [])
        if not isinstance(source_refs, list):
            raise ValueError("source_refs 必须是列表。")
        normalized_source_refs = [self._normalize_source_ref(ref) for ref in source_refs]

        tags = section_entry.get("tags", [])
        if not isinstance(tags, list):
            raise ValueError("tags 必须是列表。")
        for tag in tags:
            if not isinstance(tag, str):
                raise ValueError("tags 中的元素必须是字符串。")

        normalized = {
            "section_number": section_number,
            "section_title": section_title,
            "heading_level": heading_level,
            "content": normalized_content,
            "source_refs": normalized_source_refs,
            "tags": tags,
        }

        if "section_id" in section_entry and section_entry["section_id"] is not None:
            if not isinstance(section_entry["section_id"], str) or not section_entry["section_id"].strip():
                raise ValueError("section_id 必须是非空字符串。")
            normalized["section_id"] = section_entry["section_id"].strip()

        return normalized

    def _find_section(
        self,
        data: Dict[str, Any],
        section_number: Optional[str] = None,
        section_title: Optional[str] = None,
    ) -> Optional[Dict[str, Any]]:
        number = self._normalize_optional_str(section_number, "section_number")
        title = self._normalize_optional_str(section_title, "section_title")

        if not number and not title:
            raise ValueError("section_number 和 section_title 至少要提供一个。")

        matched = []
        for entry in data["sections"]:
            number_ok = True if not number else entry.get("section_number", "") == number
            title_ok = True if not title else entry.get("section_title", "") == title
            if number_ok and title_ok:
                matched.append(entry)

        if len(matched) > 1:
            raise ValueError("匹配到多个小节，请同时提供更精确的 section_number 和 section_title。")
        return matched[0] if matched else None

    def _ensure_unique_section(self, data: Dict[str, Any], section_entry: Dict[str, Any]) -> None:
        section_number = section_entry.get("section_number", "")
        section_title = section_entry.get("section_title", "")

        for entry in data["sections"]:
            if section_number and entry.get("section_number", "") == section_number:
                raise ValueError(f"小节编号 '{section_number}' 已存在，不能重复 add。")
            if section_title and entry.get("section_title", "") == section_title:
                raise ValueError(f"小节标题 '{section_title}' 已存在，不能重复 add。")

    # ========== 业务功能 ==========
    def set_doc_meta(self, doc_name: Optional[str], doc_path: Optional[str], doc_title: Optional[str]) -> Dict[str, Any]:
        data = self._load()

        if doc_name is not None:
            if not isinstance(doc_name, str) or not doc_name.strip():
                raise ValueError("doc_name 必须是非空字符串。")
            if not doc_name.endswith(".docx"):
                raise ValueError("doc_name 必须是 .docx 文件名。")
            data["doc_name"] = doc_name.strip()

        if doc_path is not None:
            if not isinstance(doc_path, str) or not doc_path.strip():
                raise ValueError("doc_path 必须是非空字符串。")
            if not doc_path.endswith(".docx"):
                raise ValueError("doc_path 必须指向 .docx 文件。")
            data["doc_path"] = str(Path(doc_path).expanduser().resolve())

        if doc_title is not None:
            if not isinstance(doc_title, str):
                raise ValueError("doc_title 必须是字符串。")
            data["doc_title"] = doc_title.strip()

        self._save(data)
        return {
            "success": True,
            "message": "DOC 元信息已更新。",
            "data": {
                "doc_name": data["doc_name"],
                "doc_path": data["doc_path"],
                "doc_title": data["doc_title"],
            }
        }

    def add(self, section_entry: Dict[str, Any], insert_after: Optional[Dict[str, str]] = None) -> Dict[str, Any]:
        data = self._load()
        normalized = self._normalize_section_entry(section_entry)
        self._ensure_unique_section(data, normalized)

        normalized["section_id"] = normalized.get("section_id") or self._generate_next_section_id(data)

        if insert_after is None:
            data["sections"].append(normalized)
            insert_index = len(data["sections"]) - 1
        else:
            if not isinstance(insert_after, dict):
                raise ValueError("insert_after 必须是字典。")

            anchor = self._find_section(
                data,
                section_number=insert_after.get("section_number"),
                section_title=insert_after.get("section_title"),
            )
            if anchor is None:
                raise KeyError("insert_after 指定的小节不存在。")

            anchor_index = data["sections"].index(anchor)
            data["sections"].insert(anchor_index + 1, normalized)
            insert_index = anchor_index + 1

        self._save(data)
        return {
            "success": True,
            "message": f"已新增小节 {normalized.get('section_number', '')} {normalized['section_title']}".strip(),
            "data": {
                "insert_index": insert_index,
                "section": normalized
            }
        }

    def update(
        self,
        section_number: Optional[str],
        section_title: Optional[str],
        block_name: str,
        new_content: Any,
    ) -> Dict[str, Any]:
        if block_name not in self.VALID_SECTION_BLOCKS:
            raise ValueError(f"block_name 非法，必须是 {sorted(self.VALID_SECTION_BLOCKS)} 之一。")

        data = self._load()
        entry = self._find_section(data, section_number=section_number, section_title=section_title)
        if entry is None:
            raise KeyError("未找到目标小节。")

        if block_name == "heading_level":
            if not isinstance(new_content, int) or new_content <= 0:
                raise ValueError("heading_level 的 new_content 必须是正整数。")
            entry[block_name] = new_content

        elif block_name in {"section_number", "section_title"}:
            if not isinstance(new_content, str):
                raise ValueError(f"{block_name} 的 new_content 必须是字符串。")
            new_value = new_content.strip()

            for other in data["sections"]:
                if other is entry:
                    continue
                if block_name == "section_number" and new_value and other.get("section_number", "") == new_value:
                    raise ValueError(f"小节编号 '{new_value}' 已存在。")
                if block_name == "section_title" and other.get("section_title", "") == new_value:
                    raise ValueError(f"小节标题 '{new_value}' 已存在。")

            entry[block_name] = new_value

        elif block_name == "content":
            if not isinstance(new_content, list):
                raise ValueError("content 的 new_content 必须是列表。")
            normalized_content = []
            for block in new_content:
                normalized_block = self._normalize_content_block(
                    block,
                    auto_block_id=self._generate_next_block_id(normalized_content)
                )
                if not normalized_block["block_id"]:
                    normalized_block["block_id"] = self._generate_next_block_id(normalized_content)
                normalized_content.append(normalized_block)
            entry["content"] = normalized_content

        elif block_name == "source_refs":
            if not isinstance(new_content, list):
                raise ValueError("source_refs 的 new_content 必须是列表。")
            entry["source_refs"] = [self._normalize_source_ref(ref) for ref in new_content]

        elif block_name == "tags":
            if not isinstance(new_content, list):
                raise ValueError("tags 的 new_content 必须是列表。")
            for tag in new_content:
                if not isinstance(tag, str):
                    raise ValueError("tags 中的元素必须是字符串。")
            entry["tags"] = new_content

        self._save(data)
        return {
            "success": True,
            "message": f"已更新小节 {entry.get('section_number', '')} {entry.get('section_title', '')} 的 {block_name}".strip(),
            "data": {
                "section_id": entry.get("section_id", ""),
                "section_number": entry.get("section_number", ""),
                "section_title": entry.get("section_title", ""),
                "updated_block": block_name,
                "new_content": entry.get(block_name)
            }
        }

    def delete(self, section_number: Optional[str], section_title: Optional[str]) -> Dict[str, Any]:
        data = self._load()
        entry = self._find_section(data, section_number=section_number, section_title=section_title)
        if entry is None:
            raise KeyError("未找到目标小节。")

        removed = {
            "section_id": entry.get("section_id", ""),
            "section_number": entry.get("section_number", ""),
            "section_title": entry.get("section_title", ""),
        }
        data["sections"].remove(entry)
        self._save(data)

        return {
            "success": True,
            "message": f"已删除小节 {removed['section_number']} {removed['section_title']}".strip(),
            "data": removed
        }

    def list_sections(self) -> Dict[str, Any]:
        data = self._load()
        outlines = []
        for idx, entry in enumerate(data["sections"]):
            outlines.append({
                "order": idx,
                "section_id": entry.get("section_id", ""),
                "section_number": entry.get("section_number", ""),
                "section_title": entry.get("section_title", ""),
                "heading_level": entry.get("heading_level", 2),
                "content_block_count": len(entry.get("content", [])),
            })

        return {
            "success": True,
            "message": f"当前文档共有 {len(outlines)} 个小节。",
            "data": {
                "doc_name": data.get("doc_name", ""),
                "doc_path": data.get("doc_path", ""),
                "doc_title": data.get("doc_title", ""),
                "sections": outlines
            }
        }

    def query(
        self,
        section_number: Optional[str] = None,
        section_title: Optional[str] = None,
    ) -> Dict[str, Any]:
        data = self._load()

        if (section_number is None or str(section_number).strip() == "") and \
           (section_title is None or str(section_title).strip() == ""):
            return {
                "success": True,
                "message": "已查询整个 DOC JSON。",
                "data": data
            }

        entry = self._find_section(data, section_number=section_number, section_title=section_title)
        if entry is None:
            raise KeyError("未找到目标小节。")

        return {
            "success": True,
            "message": f"已查询到小节 {entry.get('section_number', '')} {entry.get('section_title', '')}".strip(),
            "data": entry
        }


def wr_doc_json(operation: str, doc_json_path: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    """
    DOC JSON 工具统一入口。

    支持:
    - set_meta: 设置文档元信息
    - add:      新增一个小节
    - update:   更新一个小节的一个块
    - delete:   删除一个小节
    - list:     列出小节目录
    - query:    查询整个文档或某个小节
    """
    try:
        if not isinstance(operation, str) or not operation.strip():
            raise ValueError("operation 必须是非空字符串。")
        if not isinstance(doc_json_path, str) or not doc_json_path.strip():
            raise ValueError("doc_json_path 必须是非空字符串。")
        if not isinstance(payload, dict):
            raise ValueError("payload 必须是字典。")

        store = DocJsonStore(doc_json_path)
        op = operation.strip().lower()

        if op == "set_meta":
            return store.set_doc_meta(
                doc_name=payload.get("doc_name"),
                doc_path=payload.get("doc_path"),
                doc_title=payload.get("doc_title"),
            )

        elif op == "add":
            section_entry = payload.get("section_entry")
            if section_entry is None:
                raise ValueError("add 操作需要 payload['section_entry']")
            return store.add(
                section_entry=section_entry,
                insert_after=payload.get("insert_after"),
            )

        elif op in {"update", "updata"}:
            block_name = payload.get("block_name")
            if block_name is None:
                raise ValueError("update 操作需要 payload['block_name']")
            if "new_content" not in payload:
                raise ValueError("update 操作需要 payload['new_content']")

            return store.update(
                section_number=payload.get("section_number"),
                section_title=payload.get("section_title"),
                block_name=block_name,
                new_content=payload.get("new_content"),
            )

        elif op == "delete":
            return store.delete(
                section_number=payload.get("section_number"),
                section_title=payload.get("section_title"),
            )

        elif op == "list":
            return store.list_sections()

        elif op == "query":
            return store.query(
                section_number=payload.get("section_number"),
                section_title=payload.get("section_title"),
            )

        else:
            raise ValueError("operation 仅支持 set_meta / add / update / updata / delete / list / query")

    except Exception as e:
        return {
            "success": False,
            "message": f"{type(e).__name__}: {e}",
            "data": None
        }


class DocFeedbackStore:
    """
    管理 DOC 反馈 JSON 的工具类。
    支持:
    - add: 新增一个 doc 反馈条目
    - update: 更新指定 doc 的某一个块
    - list: 列出当前已有的 doc 反馈条目
    - query: 查询指定 doc 的完整反馈条目
    """

    VALID_BLOCKS = {
        "overall_assessment",
        "must_fix",
        "should_fix",
        "forbidden_content_found",
        "section_feedback",
        "recommended_structure",
    }

    def __init__(self, doc_feedback_json_path: str):
        self.db_path = Path(doc_feedback_json_path).expanduser().resolve()

    # ========== 基础读写 ==========
    def _init_db(self) -> Dict[str, Any]:
        return {
            "doc_feedbacks": []
        }

    def _load(self) -> Dict[str, Any]:
        if not self.db_path.exists():
            return self._init_db()

        try:
            text = self.db_path.read_text(encoding="utf-8").strip()
            if not text:
                data = self._init_db()
                self._save(data)
                return data

            data = json.loads(text)

            if not isinstance(data, dict):
                raise ValueError("JSON 根节点必须是字典。")
            if "doc_feedbacks" not in data or not isinstance(data["doc_feedbacks"], list):
                raise ValueError("JSON 中必须包含 'doc_feedbacks' 列表。")

            return data
        except Exception as e:
            raise ValueError(f"读取 DOC 反馈 JSON 数据库失败: {e}")

    def _save(self, data: Dict[str, Any]) -> None:
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        self.db_path.write_text(
            json.dumps(data, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )

    # ========== 辅助函数 ==========
    def _normalize_doc_name(self, doc_name: str) -> str:
        if not isinstance(doc_name, str) or not doc_name.strip():
            raise ValueError("doc_name 必须是非空字符串。")
        return doc_name.strip()

    def _find_entry_by_doc_name(self, data: Dict[str, Any], doc_name: str) -> Optional[Dict[str, Any]]:
        doc_name = self._normalize_doc_name(doc_name)
        for entry in data["doc_feedbacks"]:
            if entry.get("doc_name") == doc_name:
                return entry
        return None

    def _validate_feedback_entry(self, feedback_entry: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(feedback_entry, dict):
            raise ValueError("feedback_entry 必须是字典。")

        doc_name = feedback_entry.get("doc_name", "")
        doc_path = feedback_entry.get("doc_path", "")

        if not isinstance(doc_name, str) or not doc_name.strip():
            raise ValueError("feedback_entry['doc_name'] 必须是非空字符串。")
        if not doc_name.endswith(".docx"):
            raise ValueError("doc_name 必须是 .docx 文件。")

        if not isinstance(doc_path, str) or not doc_path.strip():
            raise ValueError("feedback_entry['doc_path'] 必须是非空字符串。")

        overall_assessment = feedback_entry.get("overall_assessment", "")
        must_fix = feedback_entry.get("must_fix", [])
        should_fix = feedback_entry.get("should_fix", [])
        forbidden_content_found = feedback_entry.get("forbidden_content_found", [])
        section_feedback = feedback_entry.get("section_feedback", [])
        recommended_structure = feedback_entry.get("recommended_structure", [])

        if not isinstance(overall_assessment, str):
            raise ValueError("overall_assessment 必须是字符串。")
        if not isinstance(must_fix, list):
            raise ValueError("must_fix 必须是列表。")
        if not isinstance(should_fix, list):
            raise ValueError("should_fix 必须是列表。")
        if not isinstance(forbidden_content_found, list):
            raise ValueError("forbidden_content_found 必须是列表。")
        if not isinstance(section_feedback, list):
            raise ValueError("section_feedback 必须是列表。")
        if not isinstance(recommended_structure, list):
            raise ValueError("recommended_structure 必须是列表。")

        normalized = {
            "doc_name": doc_name.strip(),
            "doc_path": str(Path(doc_path).expanduser().resolve()),
            "overall_assessment": overall_assessment,
            "must_fix": must_fix,
            "should_fix": should_fix,
            "forbidden_content_found": forbidden_content_found,
            "section_feedback": section_feedback,
            "recommended_structure": recommended_structure,
        }
        return normalized

    # ========== 业务功能 ==========
    def add(self, feedback_entry: Dict[str, Any]) -> Dict[str, Any]:
        """
        在 JSON 末尾追加一个新的 DOC 反馈条目。
        若 doc_name 已存在，则报错，不覆盖。
        """
        data = self._load()
        normalized = self._validate_feedback_entry(feedback_entry)

        existing = self._find_entry_by_doc_name(data, normalized["doc_name"])
        if existing is not None:
            raise ValueError(f"DOC 反馈条目 '{normalized['doc_name']}' 已存在，不能重复 add。")

        data["doc_feedbacks"].append(normalized)
        self._save(data)

        return {
            "success": True,
            "message": f"已新增 DOC 反馈条目 {normalized['doc_name']}。",
            "data": normalized
        }

    def update(self, doc_name: str, block_name: str, new_content: Any) -> Dict[str, Any]:
        """
        更新指定 doc 的某一个块。
        block_name 只能是:
        - overall_assessment
        - must_fix
        - should_fix
        - forbidden_content_found
        - section_feedback
        - recommended_structure
        """
        if block_name not in self.VALID_BLOCKS:
            raise ValueError(
                f"block_name 非法，必须是 {sorted(self.VALID_BLOCKS)} 之一。"
            )

        # 类型检查
        if block_name == "overall_assessment":
            if not isinstance(new_content, str):
                raise ValueError("overall_assessment 的 new_content 必须是字符串。")
        else:
            if not isinstance(new_content, list):
                raise ValueError(f"{block_name} 的 new_content 必须是列表。")

        data = self._load()
        entry = self._find_entry_by_doc_name(data, doc_name)
        if entry is None:
            raise KeyError(f"未找到 DOC 反馈条目: {doc_name}")

        entry[block_name] = new_content
        self._save(data)

        return {
            "success": True,
            "message": f"已更新 DOC {doc_name} 的 {block_name}。",
            "data": {
                "doc_name": entry.get("doc_name", ""),
                "updated_block": block_name,
                "new_content": entry[block_name]
            }
        }

    def list_docs(self) -> Dict[str, Any]:
        """
        列举当前 DOC 反馈 JSON 中已有的条目。
        """
        data = self._load()
        docs = []

        for entry in data["doc_feedbacks"]:
            docs.append({
                "doc_name": entry.get("doc_name", ""),
                "doc_path": entry.get("doc_path", "")
            })

        return {
            "success": True,
            "message": f"当前已有 {len(docs)} 个 DOC 反馈条目。",
            "data": docs
        }

    def query(self, doc_name: str) -> Dict[str, Any]:
        """
        查询指定 doc 的完整反馈条目。
        """
        data = self._load()
        entry = self._find_entry_by_doc_name(data, doc_name)
        if entry is None:
            raise KeyError(f"未找到 DOC 反馈条目: {doc_name}")

        return {
            "success": True,
            "message": f"已查询到 DOC {doc_name} 的反馈内容。",
            "data": entry
        }


def wr_doc_feedback_json(operation: str, doc_feedback_json_path: str, payload: Dict[str, Any]) -> Dict[str, Any]:
    """
    DOC 反馈 JSON 的统一入口函数。

    参数:
        operation:
            - "add"
            - "update"
            - "updata"
            - "list"
            - "query"

        doc_feedback_json_path:
            DOC 反馈 JSON 数据库路径

        payload:
            不同操作对应不同字段

    返回:
        dict:
        {
            "success": bool,
            "message": str,
            "data": ...
        }
    """
    try:
        store = DocFeedbackStore(doc_feedback_json_path)
        op = operation.strip().lower()

        if op == "add":
            feedback_entry = payload.get("feedback_entry")
            if feedback_entry is None:
                raise ValueError("add 操作需要 payload['feedback_entry']")
            return store.add(feedback_entry)

        elif op in {"update", "updata"}:
            doc_name = payload.get("doc_name")
            block_name = payload.get("block_name")
            new_content = payload.get("new_content")

            if doc_name is None:
                raise ValueError("update 操作需要 payload['doc_name']")
            if block_name is None:
                raise ValueError("update 操作需要 payload['block_name']")
            if new_content is None:
                raise ValueError("update 操作需要 payload['new_content']")

            return store.update(doc_name, block_name, new_content)

        elif op == "list":
            return store.list_docs()

        elif op == "query":
            doc_name = payload.get("doc_name")
            if doc_name is None:
                raise ValueError("query 操作需要 payload['doc_name']")
            return store.query(doc_name)

        else:
            raise ValueError("operation 仅支持 add / update / updata / list / query")

    except Exception as e:
        return {
            "success": False,
            "message": f"{type(e).__name__}: {e}",
            "data": None
        }

# doc-json转docx工具
def doc_json_to_docx(doc_json_path: str, output_doc_path: str = None) -> Dict[str, Any]:
    """
    将 wr_doc_json 维护的 DOC JSON 映射文件转换为 docx 文档。
    """
    try:
        json_path = Path(doc_json_path).expanduser().resolve()
        if not json_path.exists():
            raise FileNotFoundError(f"DOC JSON 文件不存在: {json_path}")

        data = json.loads(json_path.read_text(encoding="utf-8"))
        if not isinstance(data, dict):
            raise ValueError("DOC JSON 根节点必须是字典。")
        if "sections" not in data or not isinstance(data["sections"], list):
            raise ValueError("DOC JSON 中必须包含 sections 列表。")

        if output_doc_path is None:
            doc_path_str = data.get("doc_path", "")
            if isinstance(doc_path_str, str) and doc_path_str.strip():
                output_path = Path(doc_path_str).expanduser().resolve()
            else:
                output_path = json_path.with_suffix(".docx")
        else:
            output_path = Path(output_doc_path).expanduser().resolve()

        if output_path.suffix.lower() != ".docx":
            raise ValueError("输出路径必须是 .docx 文件。")

        doc = Document()

        doc_title = data.get("doc_title", "")
        if isinstance(doc_title, str) and doc_title.strip():
            doc.add_heading(doc_title.strip(), level=0)

        for section in data["sections"]:
            section_number = str(section.get("section_number", "") or "").strip()
            section_title = str(section.get("section_title", "") or "").strip()
            heading_level = int(section.get("heading_level", 2) or 2)

            heading_text = f"{section_number} {section_title}".strip()
            if heading_text:
                doc.add_heading(heading_text, level=heading_level)

            for block in section.get("content", []):
                block_type = block.get("block_type", "paragraph")
                text = str(block.get("text", ""))

                if block_type == "bullet":
                    doc.add_paragraph(text, style="List Bullet")
                else:
                    para = doc.add_paragraph(text)
                    if block_type == "code":
                        try:
                            para.style = "No Spacing"
                        except Exception:
                            pass

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))

        return {
            "success": True,
            "message": f"已成功将 DOC JSON 转换为 docx: {output_path.name}",
            "data": {
                "doc_json_path": str(json_path),
                "output_doc_path": str(output_path),
                "section_count": len(data["sections"])
            }
        }

    except Exception as e:
        return {
            "success": False,
            "message": f"{type(e).__name__}: {e}",
            "data": None
        }

if __name__ == "__main__":
    def print_divider(title: str):
        print("\n" + "=" * 20 + f" {title} " + "=" * 20)


    def assert_success(result, step_name):
        print(f"[{step_name}] 返回结果：")
        print(result)
        assert isinstance(result, dict), f"{step_name} 返回值不是 dict"
        assert result.get("success") is True, f"{step_name} 执行失败: {result}"


    # =========================
    # 1. 准备测试路径
    # =========================
    test_dir = Path("./tmp_doc_json_test").resolve()
    test_dir.mkdir(parents=True, exist_ok=True)

    doc_json_path = test_dir / "test_doc.json"
    docx_path = test_dir / "test_doc.docx"

    # 清理旧文件
    if doc_json_path.exists():
        doc_json_path.unlink()
    if docx_path.exists():
        docx_path.unlink()

    print_divider("测试路径")
    print("doc_json_path =", doc_json_path)
    print("docx_path     =", docx_path)

    # =========================
    # 2. set_meta
    # =========================
    print_divider("1. set_meta")
    result = wr_doc_json(
        operation="set_meta",
        doc_json_path=str(doc_json_path),
        payload={
            "doc_name": "test_doc.docx",
            "doc_path": str(docx_path),
            "doc_title": "固件说明文档测试"
        }
    )
    assert_success(result, "set_meta")

    # =========================
    # 3. add 第一个小节
    # =========================
    print_divider("2. add section 1")
    result = wr_doc_json(
        operation="add",
        doc_json_path=str(doc_json_path),
        payload={
            "section_entry": {
                "section_number": "1.1",
                "section_title": "时基初始化",
                "heading_level": 2,
                "content": [
                    {
                        "block_type": "paragraph",
                        "text": "系统启动后需要先建立统一时基，否则延时和定时相关功能无法可靠运行。"
                    },
                    {
                        "block_type": "code",
                        "language": "c",
                        "text": "void Stm32_Clock_Init(u8 PLL) { /* ... */ }"
                    },
                    {
                        "block_type": "paragraph",
                        "text": "该函数通过配置系统时钟和 SysTick，为后续程序提供统一的时间基准。"
                    }
                ],
                "source_refs": [
                    {
                        "file_name": "clock.v",
                        "file_path": str(test_dir / "clock.v"),
                        "line_start": 1,
                        "line_end": 25
                    }
                ],
                "tags": ["初始化", "时钟"]
            }
        }
    )
    assert_success(result, "add section 1")

    # =========================
    # 4. add 第二个小节
    # =========================
    print_divider("3. add section 2")
    result = wr_doc_json(
        operation="add",
        doc_json_path=str(doc_json_path),
        payload={
            "section_entry": {
                "section_number": "1.2",
                "section_title": "串口发送",
                "heading_level": 2,
                "content": [
                    {
                        "block_type": "paragraph",
                        "text": "为了将系统状态输出到上位机，需要实现串口发送功能。"
                    },
                    {
                        "block_type": "code",
                        "language": "c",
                        "text": "void UART_SendByte(u8 data) { /* ... */ }"
                    },
                    {
                        "block_type": "paragraph",
                        "text": "该函数通过轮询发送完成标志位，将一个字节写入串口发送寄存器。"
                    }
                ],
                "source_refs": [
                    {
                        "file_name": "uart_tx.v",
                        "file_path": str(test_dir / "uart_tx.v"),
                        "line_start": 10,
                        "line_end": 40
                    }
                ],
                "tags": ["串口", "发送"]
            }
        }
    )
    assert_success(result, "add section 2")

    # =========================
    # 5. list
    # =========================
    print_divider("4. list")
    result = wr_doc_json(
        operation="list",
        doc_json_path=str(doc_json_path),
        payload={}
    )
    assert_success(result, "list")

    sections = result["data"]["sections"]
    assert len(sections) == 2, f"list 后 section 数量错误，应为 2，实际为 {len(sections)}"

    # =========================
    # 6. query 单个小节
    # =========================
    print_divider("5. query section 1.1")
    result = wr_doc_json(
        operation="query",
        doc_json_path=str(doc_json_path),
        payload={
            "section_number": "1.1"
        }
    )
    assert_success(result, "query section 1.1")
    assert result["data"]["section_title"] == "时基初始化"

    # =========================
    # 7. update 小节内容
    # =========================
    print_divider("6. update section 1.1 content")
    result = wr_doc_json(
        operation="update",
        doc_json_path=str(doc_json_path),
        payload={
            "section_number": "1.1",
            "block_name": "content",
            "new_content": [
                {
                    "block_type": "paragraph",
                    "text": "系统启动阶段必须先建立统一时基，否则延时控制与周期调度将失去参考。"
                },
                {
                    "block_type": "code",
                    "language": "c",
                    "text": "void Stm32_Clock_Init(u8 PLL) { /* updated ... */ }"
                },
                {
                    "block_type": "paragraph",
                    "text": "更新后的描述：代码首先配置主时钟，再设置 SysTick，实现统一的微秒和毫秒时基。"
                }
            ]
        }
    )
    assert_success(result, "update section 1.1 content")

    # 再查一次，确认更新生效
    result = wr_doc_json(
        operation="query",
        doc_json_path=str(doc_json_path),
        payload={
            "section_number": "1.1"
        }
    )
    assert_success(result, "query section 1.1 after update")
    updated_texts = [block["text"] for block in result["data"]["content"]]
    assert "更新后的描述" in updated_texts[-1], "update 后内容未生效"

    # =========================
    # 8. delete 第二个小节
    # =========================
    print_divider("7. delete section 1.2")
    result = wr_doc_json(
        operation="delete",
        doc_json_path=str(doc_json_path),
        payload={
            "section_number": "1.2"
        }
    )
    assert_success(result, "delete section 1.2")

    # 再 list，确认只剩一个小节
    result = wr_doc_json(
        operation="list",
        doc_json_path=str(doc_json_path),
        payload={}
    )
    assert_success(result, "list after delete")
    sections = result["data"]["sections"]
    assert len(sections) == 1, f"delete 后 section 数量错误，应为 1，实际为 {len(sections)}"
    assert sections[0]["section_number"] == "1.1"

    # =========================
    # 9. query whole
    # =========================
    print_divider("8. query whole doc json")
    result = wr_doc_json(
        operation="query",
        doc_json_path=str(doc_json_path),
        payload={}
    )
    assert_success(result, "query whole doc json")

    whole_doc = result["data"]
    assert whole_doc["doc_title"] == "固件说明文档测试"
    assert len(whole_doc["sections"]) == 1

    # 顺便打印 JSON 内容，便于人工检查
    print("\n当前 DOC JSON 内容：")
    print(json.dumps(whole_doc, ensure_ascii=False, indent=2))

    # =========================
    # 10. 转换成 docx
    # =========================
    print_divider("9. doc_json_to_docx")
    result = doc_json_to_docx(
        doc_json_path=str(doc_json_path),
        output_doc_path=str(docx_path)
    )
    assert_success(result, "doc_json_to_docx")
    assert docx_path.exists(), "doc_json_to_docx 执行后 docx 文件未生成"

    # =========================
    # 11. 读取 docx，验证内容
    # =========================
    print_divider("10. verify generated docx")
    doc = Document(str(docx_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    print("生成的 docx 段落内容：")
    for i, p in enumerate(paragraphs, 1):
        print(f"{i}. {p}")

    # 关键断言
    assert any("固件说明文档测试" in p for p in paragraphs), "docx 中缺少文档标题"
    assert any("1.1 时基初始化" in p for p in paragraphs), "docx 中缺少小节标题"
    assert any("系统启动阶段必须先建立统一时基" in p for p in paragraphs), "docx 中缺少更新后的正文"
    assert any("void Stm32_Clock_Init" in p for p in paragraphs), "docx 中缺少代码块文本"

    print_divider("全部测试通过")
    print("wr_doc_json 和 doc_json_to_docx 基本功能验证成功。")
